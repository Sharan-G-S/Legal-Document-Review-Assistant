import pypdf
import pdfplumber
from docx import Document as DocxDocument
from pathlib import Path
from typing import Tuple, Optional

class TextExtractor:
    """Extract text from various document formats"""
    
    @staticmethod
    def extract_from_pdf(file_path: str) -> Tuple[str, int]:
        """
        Extract text from PDF file
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Tuple of (extracted_text, page_count)
        """
        text = ""
        page_count = 0
        
        try:
            # Try pdfplumber first (better for complex PDFs)
            with pdfplumber.open(file_path) as pdf:
                page_count = len(pdf.pages)
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
            
            # If pdfplumber fails or returns empty, try pypdf
            if not text.strip():
                with open(file_path, 'rb') as file:
                    pdf_reader = pypdf.PdfReader(file)
                    page_count = len(pdf_reader.pages)
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n\n"
        
        except Exception as e:
            raise Exception(f"Error extracting text from PDF: {str(e)}")
        
        return text.strip(), page_count
    
    @staticmethod
    def extract_from_docx(file_path: str) -> Tuple[str, int]:
        """
        Extract text from DOCX file
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Tuple of (extracted_text, paragraph_count)
        """
        try:
            doc = DocxDocument(file_path)
            text = ""
            paragraph_count = 0
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n\n"
                    paragraph_count += 1
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                text += "\n\n"
            
            return text.strip(), paragraph_count
        
        except Exception as e:
            raise Exception(f"Error extracting text from DOCX: {str(e)}")
    
    @staticmethod
    def extract_text(file_path: str) -> Tuple[str, int]:
        """
        Extract text from document (auto-detect format)
        
        Args:
            file_path: Path to document file
            
        Returns:
            Tuple of (extracted_text, page_or_paragraph_count)
        """
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        if extension == '.pdf':
            return TextExtractor.extract_from_pdf(str(file_path))
        elif extension in ['.docx', '.doc']:
            return TextExtractor.extract_from_docx(str(file_path))
        else:
            raise ValueError(f"Unsupported file format: {extension}")
    
    @staticmethod
    def clean_text(text: str) -> str:
        """
        Clean and normalize extracted text
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        # Remove excessive whitespace
        text = ' '.join(text.split())
        
        # Remove special characters but keep punctuation
        # text = re.sub(r'[^\w\s.,;:!?()\[\]{}"\'$%&-]', '', text)
        
        # Normalize line breaks
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        return text.strip()
